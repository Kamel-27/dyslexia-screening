import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys
import os

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_paragraph_background(p, color_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def set_paragraph_borders(p, color_hex="D1D5DB", size="12"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Left border only (like a blockquote/code border)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size) # 12 = 1.5 pt
    left.set(qn('w:space'), '10')
    left.set(qn('w:color'), color_hex)
    
    pBdr.append(left)
    pPr.append(pBdr)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Slate Blue
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(13)
    run.font.italic = True
    run.font.color.rgb = RGBColor(79, 70, 229) # Cool Indigo
    return p

def add_metadata(doc, text_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    for idx, line in enumerate(text_lines):
        run = p.add_run(line + ("\n" if idx < len(text_lines) - 1 else ""))
        run.font.name = 'Segoe UI'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(107, 114, 128) # Gray
    return p

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Slate Blue
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229) # Cool Indigo
    return p

def add_body_text(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(55, 65, 81) # Charcoal
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet_point(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        run_prefix = p.add_run(bold_prefix)
        run_prefix.font.name = 'Segoe UI'
        run_prefix.font.size = Pt(10.5)
        run_prefix.font.bold = True
        run_prefix.font.color.rgb = RGBColor(55, 65, 81)
        
    run_text = p.add_run(text)
    run_text.font.name = 'Segoe UI'
    run_text.font.size = Pt(10.5)
    run_text.font.color.rgb = RGBColor(55, 65, 81)
    return p

def add_callout(doc, text, title="IMPORTANT NOTICE"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    
    set_paragraph_background(p, "FEF3C7") # Warm Light Yellow (Amber Tint)
    set_paragraph_borders(p, "D97706", "16") # Amber left border
    
    run_title = p.add_run(f"★ {title}: ")
    run_title.font.name = 'Segoe UI'
    run_title.font.bold = True
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(180, 83, 9)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Segoe UI'
    run_text.font.italic = True
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(120, 53, 4)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    
    set_paragraph_background(p, "F3F4F6") # Slate Gray Background
    set_paragraph_borders(p, "9CA3AF", "12") # Gray border
    
    # We replace standard tabs with spaces for prettier formatting
    formatted_code = code_text.replace('\t', '    ')
    
    run = p.add_run(formatted_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 41, 55) # Off-black
    return p

def create_styled_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = False
    
    # Style header row
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_background(hdr_cells[idx], "1E3A8A") # Deep Blue
        p = hdr_cells[idx].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.runs[0]
        run.font.name = 'Segoe UI'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White
        run.font.size = Pt(10)
        
    # Style data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        # Zebra striping
        bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            row_cells[col_idx].text = str(cell_value)
            if bg_color != "FFFFFF":
                set_cell_background(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(55, 65, 81)
                
    # Apply padding and borders to cells using XML
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'bottom', 'left', 'right']:
                m = OxmlElement(f'w:{margin}')
                m.set(qn('w:w'), '120' if margin in ['left', 'right'] else '80')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
            
    # Set explicit column widths
    # For a standard 3-column table (Domain, Questions, Skills Tested)
    if len(headers) == 3:
        col_widths = [Inches(1.8), Inches(1.0), Inches(3.7)]
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(6) # Spacing after table

def main():
    doc = docx.Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Page/Section Header
    add_title(doc, "DysTest — System Logic & Pseudo-Code")
    add_subtitle(doc, "Graduation Research Project Technical Documentation")
    
    metadata = [
        "Prepared for: Graduation Project Evaluation Committee",
        "Faculty of Computing and Data Science",
        "Department of Computer Science / Data Science",
        "Date of Submission: May 2026",
        "System Version: 2.0 (Consolidated Next.js + FastAPI Platform)"
    ]
    add_metadata(doc, metadata)
    
    # Thin divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(18)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_div = p_div.add_run("__________________________________________________________________")
    run_div.font.color.rgb = RGBColor(229, 231, 235)
    
    # 1. Executive Summary
    add_heading_1(doc, "1. Executive Summary")
    add_body_text(doc, "DysTest is a standalone, interactive, gamified screening platform designed to assess the risk of dyslexia in children and adolescents aged 7–17. Inspired by the scientifically validated Dytective methodology (Rello et al., 2020), DysTest provides a non-invasive, accessible screening battery consisting of 33 gamified cognitive tasks. Crucially, the platform operates entirely in the browser without requiring specialized eye-trackers or webcams, silently recording behavioral metrics during gameplay.")
    add_body_text(doc, "The platform architecture consists of two decoupled subsystems:")
    add_bullet_point(doc, "Next.js 15 Frontend Client: Serves as the user interface, conducting the screening battery in an above-the-fold distraction-free flow. It plays automatic voice instructions and logs all user interactions (clicks, correct answers, misses, and response timing). Upon game completion, it compiles these events into comprehensive statistical features.", "1. ")
    add_bullet_point(doc, "FastAPI Machine Learning Inference Backend: Hosts three developmental age-specific Random Forest classifiers (G1: ages 7-8, G2: ages 9-11, G3: ages 12-17). It receives the feature payload, applies advanced feature corrections (clipping, adjusted accuracy for broken click mechanics), runs Random Forest inference, calculates a statistical confidence score, and returns an evaluated risk assessment.", "2. ")
    
    # 2. Gamified Cognitive Tasks & Domains
    add_heading_1(doc, "2. Cognitive Domains & Gameplay Mechanics")
    add_body_text(doc, "The screening battery comprises 33 sequential questions covering 7 crucial neuro-cognitive domains. Each domain is specifically designed to isolate and test phonological, visual, and sequential working memory skills which are frequently impaired in individuals with dyslexia. The domains are structured as follows:")
    
    headers = ["Cognitive Domain", "Questions", "Skills & Gameplay Mechanics Assessed"]
    rows = [
        ["Letter Recognition", "Q1 – Q13", "Identifies correct letters, visual/spatial rotation recognition (e.g., d vs b vs p), and mirrored glyph discrimination."],
        ["Visual Discrimination", "Q14 – Q17", "Spotting a visually distinct letter in a group (distinguishing minor structural details among graphemes)."],
        ["Phonological Awareness", "Q18 – Q25", "Auditory matching of sounds, phoneme blending, rhyme identification, and phonological pattern association."],
        ["Grapheme Mapping", "Q26 – Q27", "Evaluating sound-to-letter correspondences, grapheme replacements, and character ordering."],
        ["Syllable Processing", "Q28 – Q29", "Syllable segment recognition, re-ordering scrambled syllables, and constructing valid vocabulary words."],
        ["Working Memory", "Q30 – Q32", "Auditory-visual sequence recall, requiring participants to type back a sequential set of presented characters."],
        ["Sequence Recall", "Q33", "A final multi-element visual/auditory sequence retention and pattern reproduction task."]
    ]
    create_styled_table(doc, headers, rows)
    
    # 3. Mathematical Foundations & Feature Engineering
    add_heading_1(doc, "3. Feature Engineering & Mathematical Foundations")
    add_body_text(doc, "Raw game metrics can contain noise and anomalies due to structural gameplay differences or software constraints. DysTest V2 implements advanced data pre-processing and feature-engineering steps to guarantee robust, clean predictions and eliminate division-by-zero artifacts:")
    
    add_heading_2(doc, "3.1 Feature Value Clipping")
    add_body_text(doc, "Raw Accuracy and Missrate metrics are mathematically clamped to the closed interval [0.0, 1.0]. This correction eliminates mathematical artifacts found in original clinical datasets (where division by zero or erroneous telemetry produced values up to 875):")
    add_body_text(doc, "   ClampedValue = max( 0.0, min( 1.0, RawValue ) )", bold=True)
    
    add_heading_2(doc, "3.2 Adjusted Accuracy (Broken Gameplay Mechanics)")
    add_body_text(doc, "For certain tasks, standard accuracy calculations (Accuracy = Hits / Clicks) are structurally flawed:")
    add_bullet_point(doc, "Q26 (Letter Replacement) requires exactly 2 clicks per attempt, inflating raw Clicks.", "• ")
    add_bullet_point(doc, "Q27 & Q28 (Letter/Syllable Arrangement) involve dragging and tapping tiles. Intermediate steps and adjustments register as Clicks, artificially penalizing raw Accuracy.", "• ")
    add_bullet_point(doc, "Q30 – Q32 (Typed Sequence Recall) involve physical keystrokes. Using raw Clicks (keystrokes) as the denominator yields a meaninglessly low accuracy.", "• ")
    
    add_body_text(doc, "To correct for this, DysTest computes Adjusted Accuracy and Adjusted Missrate based on completed trial rounds rather than clicks. The denominator is formed by the sum of Hits and Misses:")
    add_body_text(doc, "   Adjusted Accuracy = Hits / ( Hits + Misses )", bold=True)
    add_body_text(doc, "   Adjusted Missrate = Misses / ( Hits + Misses )", bold=True)
    
    add_heading_2(doc, "3.3 Statistical Risk Classification & Thresholding")
    add_body_text(doc, "Three independent Random Forest models are optimized for specific developmental age groups. Instead of using a default binary classification threshold (0.50), the system implements Youden's J statistic to determine optimal thresholds, maximizing Sensitivity (true positive rate) and Specificity (true negative rate):")
    add_body_text(doc, "   Youden's J = Sensitivity + Specificity - 1", bold=True)
    add_body_text(doc, "For example, the optimal threshold for the models is determined to be approximately 0.24. If the model P(Dyslexia) >= threshold, the participant is classified with 'Dyslexia Risk'.")

    add_heading_2(doc, "3.4 Confidence Score Formula")
    add_body_text(doc, "The system measures prediction confidence by scaling the absolute distance of the probability from the point of maximum uncertainty (0.50), mapping it to a standard [0.0, 1.0] scale:")
    add_body_text(doc, "   Confidence = min( 1.0, | P(Dyslexia) - 0.50 | * 2 )", bold=True)
    add_body_text(doc, "This yields a confidence of 1.0 (100% confident) for predictions of P=0.0 or P=1.0, and 0.0 (0% confident) at P=0.50.")

    # 4. Frontend Client-Side Logic
    add_heading_1(doc, "4. Client-Side Telemetry & Scoring Engine")
    add_body_text(doc, "As the participant progresses through the 33 tasks, the Next.js frontend silently registers every action as a 'click', 'hit', or 'miss' associated with the active question ID. Upon completion, the scoring engine aggregates these events. The algorithm below documents this telemetry process:")
    
    frontend_pseudo = (
        "ALGORITHM Frontend_Session_Scoring\n"
        "INPUT:\n"
        "    Demographics: { gender: String, nativeLang: Boolean, otherLang: Boolean, age: Integer }\n"
        "    EventsList: List of SessionEvent { questionId: String, eventType: String }\n"
        "    QuestionIDs: List of Integer (1 to 33)\n\n"
        "OUTPUT:\n"
        "    ModelPayload: Key-Value Map of calculated features for API transmission\n\n"
        "BEGIN\n"
        "    // Step 1: Initialize the feature map and load demographic features\n"
        "    ModelPayload <- Empty Map\n"
        "    ModelPayload[\"Gender\"]     <- ConvertGenderToBinary(Demographics.gender) // Male = 1, Female = 0\n"
        "    ModelPayload[\"Nativelang\"] <- ConvertBooleanToBinary(Demographics.nativeLang)\n"
        "    ModelPayload[\"Otherlang\"]  <- ConvertBooleanToBinary(Demographics.otherLang)\n"
        "    ModelPayload[\"Age\"]        <- Demographics.age\n\n"
        "    // Step 2: Define question sets with non-standard typing mechanics\n"
        "    TYPED_RECALL_QUESTIONS <- {\"q30\", \"q31\", \"q32\"}\n\n"
        "    // Step 3: Compute aggregate gameplay statistics for all 33 tasks\n"
        "    FOR EACH qId IN QuestionIDs DO\n"
        "        questionKey <- Concatenate(\"q\", qId)\n"
        "        \n"
        "        // Filter events belonging to the current task\n"
        "        qEvents <- Filter EventsList where event.questionId equals questionKey\n"
        "        \n"
        "        clicks <- Count qEvents where event.eventType equals \"click\"\n"
        "        hits   <- Count qEvents where event.eventType equals \"hit\"\n"
        "        misses <- Count qEvents where event.eventType equals \"miss\"\n"
        "        score  <- hits\n"
        "        \n"
        "        // Apply custom rounding mechanics for keystroke/typed recall games\n"
        "        IF questionKey IS IN TYPED_RECALL_QUESTIONS THEN\n"
        "            rounds <- hits + misses\n"
        "            IF rounds > 0 THEN\n"
        "                accuracy <- hits / rounds\n"
        "                missrate <- misses / rounds\n"
        "            ELSE\n"
        "                accuracy <- 0.0\n"
        "                missrate <- 0.0\n"
        "            ENDIF\n"
        "        ELSE\n"
        "            IF clicks > 0 THEN\n"
        "                accuracy <- hits / clicks\n"
        "                missrate <- misses / clicks\n"
        "            ELSE\n"
        "                accuracy <- 0.0\n"
        "                missrate <- 0.0\n"
        "            ENDIF\n"
        "        ENDIF\n\n"
        "        // Map aggregated metrics into the final ML vector payload\n"
        "        ModelPayload[Concatenate(\"Clicks\", qId)]   <- clicks\n"
        "        ModelPayload[Concatenate(\"Hits\", qId)]     <- hits\n"
        "        ModelPayload[Concatenate(\"Misses\", qId)]   <- misses\n"
        "        ModelPayload[Concatenate(\"Score\", qId)]     <- score\n"
        "        ModelPayload[Concatenate(\"Accuracy\", qId)] <- accuracy\n"
        "        ModelPayload[Concatenate(\"Missrate\", qId)] <- missrate\n"
        "    ENDFOR\n\n"
        "    RETURN ModelPayload\n"
        "END"
    )
    add_code_block(doc, frontend_pseudo)
    
    # 5. Backend FastAPI ML Inference Logic
    add_heading_1(doc, "5. Server-Side ML Processing & Inference Engine")
    add_body_text(doc, "The FastAPI backend hosts the age-specific Random Forest models. It acts as an API gateway, performing rigorous input validation, executing feature-engineering fixes, routing queries based on participant age, and conducting inference. The algorithm below documents this backend lifecycle:")
    
    backend_pseudo = (
        "ALGORITHM Backend_ML_Inference\n"
        "INPUT:\n"
        "    Payload: Key-Value Map of participant features (Demographics + gameplay metrics)\n"
        "    ModelsConfig: Configuration file defining age groups, paths, thresholds, and feature names\n\n"
        "OUTPUT:\n"
        "    PredictionResult: { probability: Float, threshold: Float, prediction: String, confidence: Float, age_group: String }\n\n"
        "BEGIN\n"
        "    // Step 1: Input Validation and Schema Safeguarding\n"
        "    IF Payload does not contain \"Age\" THEN\n"
        "        RETURN ERROR \"Age must be provided\"\n"
        "    ENDIF\n"
        "    \n"
        "    age <- Integer(Payload[\"Age\"])\n"
        "    IF age < 7 OR age > 17 THEN\n"
        "        RETURN ERROR \"Age must be between 7 and 17\"\n"
        "    ENDIF\n\n"
        "    FOR EACH demographic_field IN [\"Gender\", \"Nativelang\", \"Otherlang\"] DO\n"
        "        IF Payload[demographic_field] IS NOT IN {0, 1} THEN\n"
        "            RETURN ERROR \"Demographic field must be binary (0 or 1)\"\n"
        "        ENDIF\n"
        "    ENDFOR\n\n"
        "    // Step 2: Route request to age-specific Random Forest Model Group\n"
        "    // Model Groups: G1 (7-8 years), G2 (9-11 years), G3 (12-17 years)\n"
        "    activeGroup <- NULL\n"
        "    FOR EACH group IN ModelsConfig.groups DO\n"
        "        IF group.age_min <= age AND age <= group.age_max THEN\n"
        "            activeGroup <- group\n"
        "            BREAK\n"
        "        ENDIF\n"
        "    ENDFOR\n\n"
        "    IF activeGroup IS NULL THEN\n"
        "        RETURN ERROR \"No validated ML model found for the specified age\"\n"
        "    ENDIF\n\n"
        "    // Step 3: Feature Engineering and Data Correction (V2 Pipeline)\n"
        "    enrichedPayload <- Copy(Payload)\n\n"
        "    // Fix 1: Clamp Accuracy* and Missrate* to [0.0, 1.0] to remove division-by-zero telemetry anomalies\n"
        "    FOR EACH key, value IN enrichedPayload DO\n"
        "        IF key STARTS WITH \"Accuracy\" OR key STARTS WITH \"Missrate\" THEN\n"
        "            IF value IS NOT NULL THEN\n"
        "                enrichedPayload[key] <- Clamp(Float(value), 0.0, 1.0)\n"
        "            ELSE\n"
        "                enrichedPayload[key] <- 0.0\n"
        "            ENDIF\n"
        "        ENDIF\n"
        "    ENDFOR\n\n"
        "    // Fix 2: Re-compute Adjusted Accuracies for broken-mechanic questions (Q26, Q27, Q28, Q30, Q31, Q32)\n"
        "    ADJUSTED_ACC_QUESTIONS <- {26, 27, 28, 30, 31, 32}\n"
        "    FOR EACH q IN ADJUSTED_ACC_QUESTIONS DO\n"
        "        hits   <- Float(enrichedPayload.Get(Concatenate(\"Hits\", q), 0))\n"
        "        misses <- Float(enrichedPayload.Get(Concatenate(\"Misses\", q), 0))\n"
        "        denom  <- hits + misses\n"
        "        \n"
        "        IF denom > 0 THEN\n"
        "            enrichedPayload[Concatenate(\"AdjAcc\", q)]  <- hits / denom\n"
        "            enrichedPayload[Concatenate(\"AdjMiss\", q)] <- misses / denom\n"
        "        ELSE\n"
        "            enrichedPayload[Concatenate(\"AdjAcc\", q)]  <- 0.0\n"
        "            enrichedPayload[Concatenate(\"AdjMiss\", q)] <- 0.0\n"
        "        ENDIF\n"
        "    ENDFOR\n\n"
        "    // Step 4: Construct Feature Vector matching the active model's specific signature\n"
        "    featureVector <- Empty List\n"
        "    FOR EACH featureName IN activeGroup.feature_names DO\n"
        "        featureVector.Append(Float(enrichedPayload.Get(featureName, 0.0)))\n"
        "    ENDFOR\n\n"
        "    // Convert list to a tabular Pandas DataFrame structure for Scikit-Learn\n"
        "    dataFrame <- ConvertToDataFrame([featureVector], columns=activeGroup.feature_names)\n"
        "    \n"
        "    // predict_proba yields [P(No Risk), P(Dyslexia Risk)]\n"
        "    probabilities <- activeGroup.pipeline.predict_proba(dataFrame)\n"
        "    probabilityOfRisk <- probabilities[0][1] // Grab probability of positive class (Dyslexia)\n\n"
        "    // Step 5: Risk Classification and Confidence Estimation\n"
        "    IF probabilityOfRisk >= activeGroup.threshold THEN\n"
        "        predictionLabel <- \"Dyslexia Risk\"\n"
        "    ELSE\n"
        "        predictionLabel <- \"No Risk\"\n"
        "    ENDIF\n\n"
        "    // Calculate distance from 0.50 (maximum uncertainty) and scale to [0, 1]\n"
        "    confidenceScore <- Min(1.0, AbsoluteValue(probabilityOfRisk - 0.50) * 2)\n\n"
        "    // Step 6: Return Response\n"
        "    RETURN PredictionResult {\n"
        "        probability:   probabilityOfRisk,\n"
        "        threshold:     activeGroup.threshold,\n"
        "        prediction:    predictionLabel,\n"
        "        confidence:    confidenceScore,\n"
        "        age_group:     activeGroup.name,\n"
        "        model_version: activeGroup.model_version\n"
        "    }\n"
        "END"
    )
    add_code_block(doc, backend_pseudo)
    
    # 6. Conclusion
    add_heading_1(doc, "6. Research Validation & Conclusion")
    add_body_text(doc, "By implementing Youden's J statistic-derived decision thresholds and correcting for gameplay anomalies (clipping and round-based adjusted accuracies), DysTest provides a statistically rigorous, client-safe dyslexia screening mechanism. Decoupling the client-side Next.js game execution from the FastAPI backend ensures high efficiency, security, and portability.")
    add_callout(doc, "DysTest is a statistical screening aid designed to detect the risk profile of dyslexia using gamified cognitive performance indicators. It does NOT provide a formal clinical or medical diagnosis. All screening results are probability estimates that should be interpreted by certified educational psychologists or healthcare practitioners.", "CLINICAL DISCLAIMER")
    
    # Save the file
    out_name = "DysTest_System_Logic_and_Pseudocode.docx"
    doc.save(out_name)
    print(f"SUCCESS: Report saved as {out_name}")

if __name__ == "__main__":
    main()
